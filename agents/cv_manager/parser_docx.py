"""Extract text from DOCX files using python-docx."""

from pathlib import Path

from core.logger import get_logger

log = get_logger(__name__)


def extract_text(path: str | Path) -> str:
    """Extract all paragraph and table text from a DOCX file.

    Args:
        path: Absolute or relative path to the DOCX file.

    Returns:
        Extracted text as a single string, or empty string on failure.
    """
    try:
        from docx import Document
        doc = Document(str(path))
        lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)

        result = "\n".join(lines)
        log.info("DOCX extracted: %d chars from %s", len(result), path)
        return result
    except ImportError:
        log.error("python-docx not installed. Run: pip install python-docx")
        return ""
    except Exception as e:
        log.error("DOCX extraction failed for %s: %s", path, e)
        return ""
